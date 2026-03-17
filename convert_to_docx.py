from docx import Document
from docx.shared import Pt
import os

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} does not exist.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        # Handle Bullet points
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('    - '):
            p = doc.add_paragraph(line[6:], style='List Bullet 2')
        else:
            # Normal paragraph
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    md_file = r'c:\Users\Lenovo\gitskills\PersonalAccounting\项目经历.md'
    docx_file = r'c:\Users\Lenovo\gitskills\PersonalAccounting\项目经历.docx'
    convert_md_to_docx(md_file, docx_file)
